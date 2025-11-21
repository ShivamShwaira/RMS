import os
from docx import Document

def extract_text_from_docx(file_path):
    """Extract all text from a .docx file."""
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    # # Extract text inside tables also
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             full_text.append(cell.text)

    return "\n".join(full_text).strip()


def extract_from_folder(input_folder, output_folder):
    """Extract text from all .docx files in a folder and save as .txt."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.lower().endswith(".docx"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename.replace(".docx", ".txt"))

            text = extract_text_from_docx(input_path)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)



if __name__ == "__main__":
    input_folder = "D:/WORK/RMS/Project/dataset/all_data/sample_word"
    output_folder = "D:/WORK/RMS/Project/dataset/all_data/extracted_word"

    extract_from_folder(input_folder, output_folder)

    print("Extracted Word")